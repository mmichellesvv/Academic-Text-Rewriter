#!/usr/bin/env python
# coding: utf-8

# In[2]:


import os
import fitz
from docx import Document
import requests
from typing import List, Tuple
from langdetect import detect
import time
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import argparse

# === 0. CLI Argument Parser ===
def get_args():
    parser = argparse.ArgumentParser(description="**Controlled** way of rewriting academic text from PDF or DOCX.")
    parser.add_argument("input_file", type=str, help="Path to input file (.pdf or .docx)")
    parser.add_argument("output_file", type=str, help="Path to output file (.docx)")
    parser.add_argument("--token", type=str, required=True, help="GROQ API token")
    return parser.parse_args()

# === 1. Extract text from PDF pages/chunks ===

def extract_full_text_by_chunks_from_pdf(pdf_path: str, chunk_size: int = 12) -> List[str]:
    doc = fitz.open(pdf_path)
    chunks = []

    for page in doc:
        page_text = page.get_text().strip()
        if len(page_text) < 40:
            continue

        # Розбиваємо текст сторінки на «параграфи» по двократних переносах рядка
        paragraphs = [p.strip() for p in page_text.split("\n\n") if p.strip()]

        i = 0
        carry_over = None

        while i < len(paragraphs):
            current_chunk = paragraphs[i:i + chunk_size]

            if carry_over:
                current_chunk = [carry_over] + current_chunk
                carry_over = None

            # Якщо останній параграф закінчується на ":", переносимо його в наступний chunk
            if current_chunk and current_chunk[-1].endswith(":"):
                carry_over = current_chunk.pop(-1)

            chunk_text = "\n\n".join(current_chunk)
            if chunk_text.strip():
                chunks.append(chunk_text)

            i += chunk_size

        if carry_over:
            chunks.append(carry_over)

    return chunks


# === 2. Extract text from DOCX pages/chunks ===
def extract_full_text_by_chunks_from_docx(docx_path: str, chunk_size: int = 12) -> List[str]:
    doc = Document(docx_path)
    paragraphs = [p.text.strip() for p in doc.paragraphs if len(p.text.strip()) > 20]

    chunks = []
    i = 0
    carry_over = None

    while i < len(paragraphs):
        current_chunk = paragraphs[i:i + chunk_size]

        if carry_over:
            current_chunk = [carry_over] + current_chunk
            carry_over = None

        if current_chunk and current_chunk[-1].endswith(":"):
            carry_over = current_chunk.pop(-1)

        chunk_text = "\n\n".join(current_chunk)
        if chunk_text.strip():
            chunks.append(chunk_text)

        i += chunk_size

    if carry_over:
        chunks.append(carry_over)

    return chunks

# === 3. Paraphrase using Groq API ===
def paraphrase_groq(text: str, token: str) -> str:
    try:
        tx_lang = detect(text)
    except Exception as e:
        print(f"! Cannot detect the language of the written text: {e}")
        tx_lang = "en"  # fallback

    headers = {
        "Authorization": f"Bearer {token}",
        "Content-Type": "application/json"
    }

    payload = {
        "model": "llama3-70b-8192",
        "messages": [
            {
                "role": "system",
                "content": (
                    f"Ти експерт у сфері. Застосуй наступні рекомендації до **кожного** абзацу: "
                    f"Перебудуй речення або їх частини — так, щоб порядок слів був іншим, але сенс залишився тим самим. "
                    f"Перепиши текст іншими словами, змінюючи лише 10–20% формулювань, але повністю зберігаючи зміст. "
                    f"Додай 5% нової інформації або деталей, що збагачують текст. "
                    f"Збережи академічно-офіційний стиль. Відповідь має бути {tx_lang} мовою. "
                    f"Поверни лише перефразований текст без вступу або пояснень."
                )
            },
            {"role": "user", "content": text}
        ],
        "temperature": 0.7,
        "frequency_penalty": 0.4,
        "presence_penalty": 0.2
    }

    try:
        response = requests.post("https://api.groq.com/openai/v1/chat/completions", headers=headers, json=payload)
        response.raise_for_status()
        result = response.json()
        return result["choices"][0]["message"]["content"].strip()
    except Exception as e:
        print(f"! Paraphrasing error: {e}")
        return "[Paraphrasing Error]"

# === 4. Save to DOCX ===
def write_to_docx_by_page(pages: List[Tuple[str, str]], output_path: str):
    doc = Document()
    doc.add_heading("Academic Rewriting", 0)

    for idx, (original, paraphrased) in enumerate(pages, 1):
        heading = doc.add_heading(f"P{idx}\n", level=1)

        run = heading.runs[0]
        font = run.font
        font.name = 'Arial'           
        font.size = Pt(16)            
        font.bold = True             
        font.color.rgb = RGBColor(0x00, 0x33, 0x66)  
        heading.alignment = WD_ALIGN_PARAGRAPH.LEFT 

        # Додаємо параграф з перефразованим текстом
        doc.add_paragraph(paraphrased)
        doc.add_paragraph("")

    doc.save(output_path)

# === 5. Main Process ===
def process_document(input_path: str, output_path: str, token: str):
    ext = os.path.splitext(input_path)[1].lower()
    if ext == ".pdf":
        pages = extract_full_text_by_chunks_from_pdf(input_path)
    elif ext == ".docx":
        pages = extract_full_text_by_chunks_from_docx(input_path)
    else:
        raise ValueError("Only .pdf and .docx formats are supported!")

    print(f"\n- Processing {len(pages)} pages/chunks...")
    final_result = []

    for i, text in enumerate(pages):
        print(f"\n📄 Page {i + 1}")
        paraphrased = paraphrase_groq(text, token)
        final_result.append((text, paraphrased))
        time.sleep(0.7)

    write_to_docx_by_page(final_result, output_path)
    print(f"\n- Done successfuly! Head to: {output_path}")

# === 6. Entry Point ===
if __name__ == "__main__":
    args = get_args()
    process_document(args.input_file, args.output_file, args.token)


# In[ ]:




